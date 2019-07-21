# -*- coding: utf-8 -*-
"""
Created on Tue Feb 26 01:05:54 2019

@author: shubh
"""

import xlrd 
import matplotlib.pyplot as plt  
from docx import Document
from docx.shared import Inches
# Give the location of the file 
loc = ("Disease.xls") 
# To open Workbook 
wb = xlrd.open_workbook(loc) 
sheet = wb.sheet_by_index(0) 
  
#GOOGLE_APPLICATION_CREDENTIALS='C:\tensorflow1\PlantDiseasePrediction\google-services.json'


# For row 0 and column 0 
str1=sheet.cell_value(1, 1)
str2=sheet.cell_value(1, 2)
str3=sheet.cell_value(1, 3)

max=0
found=0

def find_all(x,col):
    ctr=0
    for i in range(sheet.nrows):
        if i==0:
            continue
        else:
            str1=sheet.cell_value(i, col)
            if str1==x:
                ctr+=1
            
                
    return ctr

def find_before(i,j,x):
    ctr=0
    for k in range(i):
        str1=sheet.cell_value(k, j)
        #print (str1)
        if str1==x:
            ctr=1
            break
                                
    return ctr

max_name=[]
max_freq=[]

for j in range(sheet.ncols):
        if j==0:
            continue
        max=-1
        for i in range(sheet.nrows):
            if i==0:
                continue
            else:
                str1=sheet.cell_value(i, j)
                #print(str1)
                p=find_before(i,j,str1)
                #print(p)
                if p==0:
                    f=find_all(str1,j)
                    if f>max:
                        max=f
                        f1=str1
        #print(f1,max)
        max_name.append(f1)
        max_freq.append(max)                

#print(max_name)
#print(max_freq)
        
        

# Data to plot
labels = max_name[0], max_name[1], max_name[2], max_name[3],max_name[4]
sizes = [max_freq[0],max_freq[1],max_freq[2],max_freq[3],max_freq[4]]
colors = ['gold', 'yellowgreen', 'lightcoral', 'lightskyblue','green']
explode = (0.1, 0, 0, 0, 0)  # explode 1st slice
 
# Plot
plt.pie(sizes, explode=explode, labels=labels, colors=colors,
autopct='%1.1f%%', shadow=True, startangle=140)
 
plt.axis('equal')
plt.savefig('Pie_chart.png')   # save the figure to file
plt.show()

#writing to wordfile
document = Document()

document.add_heading('Disease Prediction Report', 0)

p = document.add_paragraph('This is the report generated from the analysis of video captured by the drone ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

document.add_heading('Statistical analysis', level=1)
document.add_paragraph('The below graph shows the distribution of various disease as per our prediction. The prediction is done on the video captured from the drone', style='Intense Quote')
document.add_picture('Pie_chart.png' , width=Inches(4.25))

#document.add_paragraph('Intense quote', style='Intense Quote')

#document.add_paragraph('first item in unordered list', style='List Bullet')
#document.add_paragraph('first item in ordered list', style='List Number')

document.add_heading('Most Severe Disease', level=1)
document.add_paragraph('The two most severe diseases are:')
document.add_paragraph(max_name[0], style='List Number')
document.add_paragraph(max_name[1], style='List Number')

disease_name1=''
disease_name2=''
if max_name[0]=='corn maize cercospora leaf spot gray leaf spot':
    disease_name1='The fungus that causes gray leaf spot is able to survive on residue for more than one year, and economically damaging disease levels have been observed in Indiana fields with two-year-old corn residue. Fungicides are available for in-season gray leaf spot management.'
elif max_name[0]=='dfsa':
    disease_name1='dfsd'     

if max_name[1]=='dfs':
    disease_name2='dfg'
elif max_name[1]=='dfsa':
    disease_name2='dfsd'


document.add_heading('Cure Of Severe Diseases Found', level=1)
document.add_paragraph('The cure of disease are:')
document.add_paragraph(max_name[0], style='List Bullet')
document.add_paragraph(disease_name1)
document.add_paragraph(disease_name2)


document.add_heading('Detailed Analysis', level=1)
document.add_paragraph('Below is the indepth report')

new1=()
records1=()
for i in range(sheet.nrows):
        if i==0:
            continue
        l=[]
        for j in range(sheet.ncols-3):
            str1=sheet.cell_value(i, j)
            l.append(str1)
            #print(str1)
        new1=tuple(l)
        records1=records1+(new1,)
                        
#print(records1)
                        
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Disease1'
hdr_cells[2].text = 'Disease2'
for qty, id, desc in records1:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('Disease_Report.docx')

# Import gcloud
from google.cloud import storage

    # Enable Storage
from firebase import firebase

import pyrebase

config = {
    "apiKey": " AIzaSyD_aONTpekw50_3RlNtt-LApnVuO2kisbM ",
    "authDomain": "plant-7de21.firebaseapp.com",
    "databaseURL": "https://plant-7de21.firebaseio.com/",
    #"projectId": "",
    "storageBucket": "plant-7de21.appspot.com",
}

firebase = pyrebase.initialize_app(config)
storage = firebase.storage()
storage.child("Disease\Disease_Report.docx").put("Disease_Report.docx")

#firebase_admin.App("Plant",cred,"Default")
#firebase_admin.initialize_app(cred)   
#firebase_admin.initialize_app(credential=None, options=None, name='[DEFAULT]')
#firebase_admin.get_app("Plant")


    # Reference an existing bucket.


    # Upload a local file to a new file to be created in your bucket.
#zebraBlob = bucket.get_blob('zebra.jpg')
#zebraBlob.upload_from_filename(filename='/photos/zoo/zebra.jpg')

    # Download a file from your bucket.
#giraffeBlob = bucket.get_blob('giraffe.jpg')
#giraffeBlob.download_as_string()

                            